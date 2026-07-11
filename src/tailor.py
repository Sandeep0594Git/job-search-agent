"""Analyze a job description with Gemini (free tier) and generate a tailored
resume .docx from the user's baseline resume."""
import base64
import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .config import ROOT, Secrets

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)   # dark navy for name + section headings
MUTED = RGBColor(0x55, 0x55, 0x55)    # grey for contact/company lines
_DATE_RE = re.compile(
    r"\b(19|20)\d{2}\b|\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)"
    r"[a-z]*\.?\s+\d{4}", re.IGNORECASE)

log = logging.getLogger("agent.tailor")

RESUME_EXTS = [".pdf", ".docx", ".md", ".txt"]

PROMPT = """You are an expert resume writer and career coach.

Below is my BASELINE RESUME and a TARGET JOB DESCRIPTION.

Produce THREE sections in your reply, exactly in this format:

===FIT ANALYSIS===
A short analysis (max 150 words): fit score out of 10, top 3 strengths to
emphasize, top 2 gaps and how to address them, and 3 likely interview themes.

===TAILORED RESUME===
The full tailored resume, rewritten to target this job. Rules:
- Keep it truthful: only reorder, reword, emphasize and quantify what exists
  in the baseline. Never invent employers, titles, dates or credentials.
- Mirror the job description's key terminology naturally (for ATS matching).
- Lead with the most relevant experience and skills for THIS role.
- Use this plain structure: lines starting with '# ' are the candidate name,
  '## ' are section headings, '- ' are bullet points, everything else is a
  normal paragraph. No markdown bold/italic/tables.
- Keep it to roughly 1-2 pages of content.

===COVER LETTER===
A concise cover letter (max 250 words) to the hiring team at the target
company for this specific role. Truthful, specific, grounded in the baseline
resume; open with why this role, close with a confident call to action.
Plain paragraphs only, no addresses or date lines.

===BASELINE RESUME===
{resume}

===TARGET JOB DESCRIPTION===
Title: {title}
Company: {company}
{description}
"""


# ---------------- Baseline resume loading ----------------

def _extract_text(path: Path) -> str:
    if path.suffix.lower() == ".pdf":
        from pypdf import PdfReader
        return "\n".join(page.extract_text() or "" for page in PdfReader(str(path)).pages)
    if path.suffix.lower() == ".docx":
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
    return path.read_text(encoding="utf-8", errors="replace")


def load_baseline_resume(cfg: dict) -> str:
    resume_dir = ROOT / (cfg.get("tailor", {}).get("resume_dir") or "resume")
    if resume_dir.exists():
        for ext in RESUME_EXTS:
            for f in sorted(resume_dir.glob(f"*{ext}")):
                text = _extract_text(f).strip()
                if len(text) > 200:
                    log.info("baseline resume: %s (%d chars)", f.name, len(text))
                    return text

    # Fallback for public repos: resume supplied as a base64 secret.
    if Secrets.RESUME_B64:
        tmp = ROOT / "data" / f"_baseline.{Secrets.RESUME_EXT.lstrip('.')}"
        tmp.parent.mkdir(exist_ok=True)
        tmp.write_bytes(base64.b64decode(Secrets.RESUME_B64))
        return _extract_text(tmp).strip()

    raise FileNotFoundError(
        "No baseline resume found. Put your resume (.pdf/.docx/.md/.txt) in the "
        "resume/ folder, or set the RESUME_B64 secret."
    )


# ---------------- Gemini ----------------

def _gemini(prompt: str, model: str) -> str:
    from google import genai
    client = genai.Client(api_key=Secrets.GEMINI_API_KEY)
    resp = client.models.generate_content(model=model, contents=prompt)
    return resp.text or ""


# ---------------- DOCX rendering ----------------

def _bottom_border(paragraph):
    """Add a thin accent rule under a paragraph (used for section headings)."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "1F3A5F")
    pbdr.append(bottom)
    pPr.append(pbdr)


def _render_docx(resume_text: str, out_path: Path):
    """Render the plain-marked resume text into a clean, professional .docx.

    Marks produced by the model:
      '# '  -> candidate name (large, centered)
      '## ' -> section heading (accent, ruled)
      '- '  -> bullet point
      other -> paragraph; contact/tagline lines (before first section) are
               centered grey; role lines with dates are bolded; 'A | B'
               company/location lines are italic grey.
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.5)
        section.left_margin = section.right_margin = Inches(0.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.05

    in_header = False  # between the name and the first '## ' section

    for raw in resume_text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line[2:].strip().upper())
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = ACCENT
            in_header = True

        elif line.startswith("## "):
            in_header = False
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(9)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line[3:].strip().upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = ACCENT
            _bottom_border(p)

        elif re.match(r"^\s*[-*•]\s+", line):
            in_header = False
            p = doc.add_paragraph(re.sub(r"^\s*[-*•]\s+", "", line),
                                  style="List Bullet")
            p.paragraph_format.space_after = Pt(2)

        elif in_header:
            # tagline + contact block under the name
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(9.5)
            run.font.color.rgb = MUTED

        elif " | " in line:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.italic = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = MUTED

        elif _DATE_RE.search(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(10.5)

        else:
            doc.add_paragraph(line)

    out_path.parent.mkdir(exist_ok=True)
    doc.save(str(out_path))


# ---------------- Public API ----------------

def tailor_resume(job: dict, cfg: dict) -> tuple:
    """Returns (fit_analysis_text, resume_docx_path, cover_letter_docx_path).
    cover_letter_docx_path is None if the model omitted that section."""
    if not Secrets.gemini_ready():
        raise RuntimeError("GEMINI_API_KEY is not set - cannot tailor resume.")

    baseline = load_baseline_resume(cfg)
    tcfg = cfg.get("tailor") or {}
    description = job.get("description") or "(No description captured; tailor from the title and company context.)"

    reply = _gemini(
        PROMPT.format(resume=baseline[:30000], title=job["title"],
                      company=job["company"], description=description[:20000]),
        model=tcfg.get("model", "gemini-2.5-flash"),
    )

    analysis, resume_out, cover_out = "", reply, ""
    if "===TAILORED RESUME===" in reply:
        analysis, rest = reply.split("===TAILORED RESUME===", 1)
        analysis = analysis.replace("===FIT ANALYSIS===", "").strip()
        if "===COVER LETTER===" in rest:
            resume_out, cover_out = rest.split("===COVER LETTER===", 1)
        else:
            resume_out = rest
        resume_out, cover_out = resume_out.strip(), cover_out.strip()

    safe_company = re.sub(r"[^A-Za-z0-9]+", "_", job["company"])[:30] or "company"
    out_dir = ROOT / (tcfg.get("output_dir") or "output")
    resume_path = out_dir / f"Resume_{safe_company}_{job['id']}.docx"
    _render_docx(resume_out, resume_path)

    cover_path = None
    if cover_out:
        cover_path = out_dir / f"CoverLetter_{safe_company}_{job['id']}.docx"
        _render_docx(cover_out, cover_path)
    return analysis, resume_path, cover_path
